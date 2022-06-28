# -*- coding: utf-8 -*-

import re
from docx import Document
from odoo.exceptions import AccessError, UserError, ValidationError



def crear_documento_reserva(ruta,detalle,lista_estado_cuenta):
    doc = Document(ruta)
    tabla=doc.tables[1]
    contador=1
    suma=0
    for estado_cuenta in lista_estado_cuenta:
        for l in estado_cuenta:
            if l.numero_cuota!=False:
                tabla.cell(contador, 0).text = str(l.numero_cuota)
            if l['fecha']!=False:
                tabla.cell(contador, 1).text = str(l.fecha)
            if l['cuota_capital']!=False:
                tabla.cell(contador, 2).text = str(l.cuota_capital)
            if l['cuota_adm']!=False:
                tabla.cell(contador, 3).text = str(l.cuota_adm)
            if l['iva_adm']!=False:
                tabla.cell(contador, 4).text = str(l.iva_adm)
            if l['saldo']!=False:
                suma+=l.saldo
                tabla.cell(contador, 5).text = str(l.saldo)
            else:
                tabla.cell(contador, 5).text = str(0)

            contador+=1
    detalle.append({'identificar_docx':'total_deuda','valor':str(round(suma,2))})
    for campo in detalle:
        regex1 = re.compile(campo['identificar_docx'])
        docx_replace_regex_ram(doc,regex1,campo['valor'] or "" )
        docx_replace_regex_header_ram(doc.sections[0].header,regex1,campo['valor'] or "")
    doc.save(ruta)

def identificar_parrafo(doc_obj, regex):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            return p

def docx_replace_regex_header_ram(doc_obj, regex , replace):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text
    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex_ram(cell, regex , replace)

def docx_replace_regex_ram(doc_obj, regex , replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex_ram(cell, regex , replace)
