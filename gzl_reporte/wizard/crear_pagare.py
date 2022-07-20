# -*- coding: utf-8 -*-

import re
from docx import Document
from odoo.exceptions import AccessError, UserError, ValidationError


def convierte_cifra(numero,sw):
    lista_centana = ["",("CIEN","CIENTO"),"DOSCIENTOS","TRESCIENTOS","CUATROCIENTOS","QUINIENTOS","SEISCIENTOS","SETECIENTOS","OCHOCIENTOS","NOVECIENTOS"]
    lista_decena = ["",("DIEZ","ONCE","DOCE","TRECE","CATORCE","QUINCE","DIECISEIS","DIECISIETE","DIECIOCHO","DIECINUEVE"),
                    ("VEINTE","VEINTI"),("TREINTA","TREINTA Y "),("CUARENTA" , "CUARENTA Y "),
                    ("CINCUENTA" , "CINCUENTA Y "),("SESENTA" , "SESENTA Y "),
                    ("SETENTA" , "SETENTA Y "),("OCHENTA" , "OCHENTA Y "),
                    ("NOVENTA" , "NOVENTA Y ")
                ]
    lista_unidad = ["",("UN" , "UNO"),"DOS","TRES","CUATRO","CINCO","SEIS","SIETE","OCHO","NUEVE"]
    centena = int (numero / 100)
    decena = int((numero -(centena * 100))/10)
    unidad = int(numero - (centena * 100 + decena * 10))
    #print "centena: ",centena, "decena: ",decena,'unidad: ',unidad
 
    texto_centena = ""
    texto_decena = ""
    texto_unidad = ""
 
    #Validad las centenas
    texto_centena = lista_centana[centena]
    if centena == 1:
        if (decena + unidad)!=0:
            texto_centena = texto_centena[1]
        else :
            texto_centena = texto_centena[0]
 
    #Valida las decenas
    texto_decena = lista_decena[decena]
    if decena == 1 :
         texto_decena = texto_decena[unidad]
    elif decena > 1 :
        if unidad != 0 :
            texto_decena = texto_decena[1]
        else:
            texto_decena = texto_decena[0]
    #Validar las unidades
    #print "texto_unidad: ",texto_unidad
    if decena != 1:
        texto_unidad = lista_unidad[unidad]
        if unidad == 1:
            texto_unidad = texto_unidad[sw]
 
    return "%s %s %s" %(texto_centena,texto_decena,texto_unidad)

def numero_to_letras(numero):
    indicador = [("",""),("MIL","MIL"),("MILLON","MILLONES"),("MIL","MIL"),("BILLON","BILLONES")]
    entero = int(numero)
    decimal = int(round((numero - entero)*100))
    #print 'decimal : ',decimal 
    contador = 0
    numero_letras = ""
    while entero >0:
        a = entero % 1000
        if contador == 0:
            en_letras = convierte_cifra(a,1).strip()
        else :
            en_letras = convierte_cifra(a,0).strip()
        if a==0:
            numero_letras = en_letras+" "+numero_letras
        elif a==1:
            if contador in (1,3):
                numero_letras = indicador[contador][0]+" "+numero_letras
            else:
                numero_letras = en_letras+" "+indicador[contador][0]+" "+numero_letras
        else:
            numero_letras = en_letras+" "+indicador[contador][1]+" "+numero_letras
        numero_letras = numero_letras.strip()
        contador = contador + 1
        entero = int(entero / 1000)
    numero_letras = numero_letras+" con " + str(decimal) +"/100"

    return numero_letras

def crear_pagare(ruta,detalle,lista_estado_cuenta):
    doc = Document(ruta)
    tabla=doc.tables[0]
    contador=1
    suma=0
    for estado_cuenta in lista_estado_cuenta:
        for l in estado_cuenta:
            if l.saldo:
                if l.numero_cuota!=False:
                    tabla.cell(contador, 0).text = str(l.numero_cuota)
                if l['fecha']!=False:
                    tabla.cell(contador, 1).text = str(l.fecha)
                if l['cuota_capital']!=False:
                    tabla.cell(contador, 2).text = str(l.cuota_capital)
                if l['cuota_adm']!=False:
                    tabla.cell(contador, 3).text = str(l.cuota_adm)
                if l['iva_adm']!=False:
                    tabla.cell(contador, 4).text = str(l.iva_adm)
                if l['saldo']!=False:
                    suma+=l.saldo
                    tabla.cell(contador, 5).text = str(l.saldo)

                contador+=1
    detalle.append({'identificar_docx':'monto_pendiente','valor':str(round(suma,2))})
    if suma:
        valor_letras=numero_to_letras(suma)
        detalle.append({'identificar_docx':'letras_valor','valor':str(valor_letras)})
    for campo in detalle:
        regex1 = re.compile(campo['identificar_docx'])
        docx_replace_regex_ram(doc,regex1,campo['valor'] or "" )
        docx_replace_regex_header_ram(doc.sections[0].header,regex1,campo['valor'] or "")
    doc.save(ruta)

def identificar_parrafo(doc_obj, regex):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            return p

def docx_replace_regex_header_ram(doc_obj, regex , replace):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text
    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex_ram(cell, regex , replace)

def docx_replace_regex_ram(doc_obj, regex , replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex_ram(cell, regex , replace)
