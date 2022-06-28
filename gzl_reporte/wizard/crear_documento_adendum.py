# -*- coding: utf-8 -*-

import re
from docx import Document
from odoo.exceptions import AccessError, UserError, ValidationError


def crear_documento_adendum(ruta,detalle,lista_estado_cuenta,estado_cuenta_anterior):
    #Se abre el documento en la ruta
    doc = Document(ruta)
    tabla2=doc.tables[0]
    contador=1
    for estado_cuenta2 in estado_cuenta_anterior:
        for a in estado_cuenta2:
            if a['monto_financiamiento_anterior']!=False:
                tabla2.cell(contador, 1).text = '$ '+str(a.monto_financiamiento_anterior)
            if a['plazo_meses_anterior']!=False:
                tabla2.cell(contador, 3).text = str(a.plazo_meses_anterior.numero)+' Meses'
            
            #if l['cuota_capital']!=False:
            #    raise ValidationError(str(l['cuota_capital']) )
            tabla2.cell(contador, 5).text = '$ '+str(a.cuota_capital_anterior) or '$0.00'
            
            contador+=1
    #contador=1
    # # #Tabla de Alarmas Rojas
    tabla=doc.tables[1]
    contador=1
    for estado_cuenta in estado_cuenta_anterior:
        for l in estado_cuenta:
            if l['monto_financiamiento']!=False:
                tabla.cell(contador, 1).text = '$ '+str(l.monto_financiamiento)
            if l['plazo_meses']!=False:
                tabla.cell(contador, 3).text = str(l.plazo_meses.numero)+' Meses'
            
            #if l['cuota_capital']!=False:
            #    raise ValidationError(str(l['cuota_capital']) )
            tabla.cell(contador, 5).text = '$ '+str(l.cuota_capital) or '$0.00'
            
            contador+=1

    #for alarma in dct_final['detalle_carac_rojo']:
    #    if alarma['item']!=False:
    #        tabla.cell(contador, 0).text = alarma['item']
    #    if alarma['name']!=False:
    #        tabla.cell(contador, 1).text = alarma['name']
    #    if alarma['estado']!=False:
    #        tabla.cell(contador, 2).text = alarma['estado']
    #    contador+=1
        #if contador!=len(dct_final['detalle_carac_rojo'])+1:
        #    tabla.add_row()  

    for campo in detalle:
        #Redenriza
        regex1 = re.compile(campo['identificar_docx'])

      #  Reemplaza los valores de identificadores de la plantilla con los del json
        docx_replace_regex_ram(doc,regex1,campo['valor'])
        docx_replace_regex_header_ram(doc.sections[0].header,regex1,campo['valor'])


    doc.save(ruta)


def identificar_parrafo(doc_obj, regex):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            return p






def docx_replace_regex_header_ram(doc_obj, regex , replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):

            inline = p.runs


            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):

                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex_ram(cell, regex , replace)









def docx_replace_regex_ram(doc_obj, regex , replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs

            # Loop added  to work with runs (strings with same style) 
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex_ram(cell, regex , replace)

