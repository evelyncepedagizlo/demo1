# -*- coding: utf-8 -*-

import re
from docx import Document
from odoo.exceptions import AccessError, UserError, ValidationError


def crear_carta_finalizacion(ruta,detalle):
    #Se abre el documento en la ruta
    doc = Document(ruta)
    contador=1
    for campo in detalle:
        regex1 = re.compile(campo['identificar_docx'])
        docx_replace_regex_ram(doc,regex1,campo['valor'])
        docx_replace_regex_header_ram(doc.sections[0].header,regex1,campo['valor'])
    doc.save(ruta)

def identificar_parrafo(doc_obj, regex):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            return p

def docx_replace_regex_header_ram(doc_obj, regex , replace):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text
    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex_ram(cell, regex , replace)

def docx_replace_regex_ram(doc_obj, regex , replace):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex_ram(cell, regex , replace)

