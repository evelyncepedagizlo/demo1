# -*- coding: utf-8 -*-

import re
from docx import Document






def crear_reporte_entregable_hito(ruta,detalle,detalle_caracteristicas):
    #Se abre el documento en la ruta
    doc = Document(ruta)

    #Lee el json 
    for campo in detalle:
        #Redenriza
        regex1 = re.compile(campo['identificar_docx'])

        #Reemplaza los valores de identificadores de la plantilla con los del json
        docx_replace_regex(doc,regex1,campo['valor'])
        docx_replace_regex_header(doc.sections[0].header,regex1,campo['valor'])


#####Detalle de Actividades realizadas
    tabla_actividades_realizada=doc.tables[3]



    numero_filas=1
    contador_caracteristica=1

    for caracteristicas in detalle_caracteristicas:

        
        tabla_actividades_realizada.cell(numero_filas, 0).text = caracteristicas['item'] +' ' +caracteristicas['name']
        tabla_actividades_realizada.cell(numero_filas, 0).merge(tabla_actividades_realizada.cell(numero_filas, 3))
        tabla_actividades_realizada.add_row() 
        numero_filas+=1

        contador_historias=1

        for historia in caracteristicas['detalle_historias']:

            tabla_actividades_realizada.cell(numero_filas, 0).text = historia['item']
            tabla_actividades_realizada.cell(numero_filas, 1).text = historia['name']
            tabla_actividades_realizada.cell(numero_filas, 2).text = historia['estado']
            tabla_actividades_realizada.cell(numero_filas, 3).text = historia['observaciones']

            numero_filas+=1

            if contador_historias!=len( caracteristicas['detalle_historias']):
                tabla_actividades_realizada.add_row() 
            elif contador_caracteristica!=len(detalle_caracteristicas):
                tabla_actividades_realizada.add_row() 
            

            contador_historias=1


        contador_caracteristica+=1





    doc.save(ruta)



def docx_replace_regex_header(doc_obj, regex , replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)









def docx_replace_regex(doc_obj, regex , replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)

